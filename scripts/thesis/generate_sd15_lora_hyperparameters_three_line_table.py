from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path("docs/thesis/sd15-electric-specialized-lora-hyperparameters-three-line-table.docx")

ROWS = [
    ("基础模型", "runwayml/stable-diffusion-v1-5", "以 Stable Diffusion v1.5 作为 LoRA 微调底座模型。"),
    ("输出模型名", "sd15-electric-specialized", "训练完成后合并生成的电力领域专用模型名称。"),
    ("训练样本数", "40", "使用整理后的电力图像数据样本数。"),
    ("输入分辨率", "512", "训练图像统一缩放到 512×512 分辨率。"),
    ("LoRA Rank", "32", "低秩适配器秩参数，用于控制可训练增量矩阵规模。"),
    ("LoRA Alpha", "32", "LoRA 缩放系数，与 Rank 保持一致。"),
    ("单卡批大小", "1", "每次前向传播使用 1 张图像样本。"),
    ("梯度累积步数", "8", "通过梯度累积提升等效批大小。"),
    ("等效总批大小", "8", "由单卡批大小 1 和梯度累积步数 8 共同形成。"),
    ("学习率", "5e-05", "LoRA 训练阶段采用的初始学习率。"),
    ("学习率调度器", "cosine", "使用余弦退火策略调节训练过程中的学习率。"),
    ("预热步数", "200", "训练初期进行 200 步学习率预热。"),
    ("训练轮数", "100", "全量数据集共训练 100 个 epoch。"),
    ("总优化步数", "500", "日志记录的总优化更新步数。"),
    ("检查点保存间隔", "500", "每 500 步保存一次训练检查点。"),
    ("混合精度", "no", "训练时未启用混合精度。"),
    ("梯度检查点", "true", "启用梯度检查点以降低显存占用。"),
    ("中心裁剪", "true", "训练前对图像执行中心裁剪。"),
    ("随机翻转", "true", "训练时启用随机翻转增强。"),
    ("数据加载线程数", "0", "数据加载采用单进程方式。"),
    ("随机种子", "42", "用于保证实验复现性。"),
]


def set_run_font(run, east_asia_font: str, latin_font: str, size_pt: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)


def set_paragraph_style(paragraph, *, align: int, size_pt: float, bold: bool = False) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", size_pt, bold=bold)


def set_cell_borders(cell, *, top: str = "nil", bottom: str = "nil", left: str = "nil", right: str = "nil") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        if value != "nil":
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def fill_cell(cell, text: str, *, align: int, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(text)
    set_paragraph_style(paragraph, align=align, size_pt=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title = doc.add_paragraph()
    title.add_run("sd15-electric-specialized LoRA训练关键超参数")
    set_paragraph_style(title, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=18, bold=True)

    doc.add_paragraph()

    caption = doc.add_paragraph()
    caption.add_run("表4.1 sd15-electric-specialized LoRA训练关键超参数")
    set_paragraph_style(caption, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)

    table = doc.add_table(rows=1 + len(ROWS), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(4.0), Cm(4.0), Cm(8.0)]
    headers = ["参数项", "取值", "说明"]

    for column, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[column]
        cell.width = width
        fill_cell(cell, header, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        set_cell_borders(cell, top="single", bottom="single")

    for row_index, values in enumerate(ROWS, start=1):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.85)
        for column, (text, width) in enumerate(zip(values, widths)):
            cell = row.cells[column]
            cell.width = width
            align = WD_ALIGN_PARAGRAPH.LEFT if column == 2 else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(cell, text, align=align)
            is_last_row = row_index == len(ROWS)
            set_cell_borders(cell, bottom="single" if is_last_row else "nil")

    return doc


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_doc()
    document.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
