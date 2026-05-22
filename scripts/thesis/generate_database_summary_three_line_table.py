from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT = Path("docs/thesis/electric-ai-platform-database-summary-three-line-table.docx")

ROWS = [
    (
        "1",
        "auth_users",
        "用户认证表",
        "用于保存平台登录用户的账号、密码摘要、显示名称和账户状态等基础认证信息。",
    ),
    (
        "2",
        "model_registry",
        "模型注册表",
        "用于维护生成模型与评分模型的名称、类型、服务归属、默认提示词和启用状态等元数据。",
    ),
    (
        "3",
        "model_prompt_templates",
        "提示词模板表",
        "用于保存不同业务场景下的正向提示词与反向提示词模板，便于快速复用。",
    ),
    (
        "4",
        "task_jobs",
        "任务主表",
        "用于统一记录图像生成与评分任务的类型、状态、执行阶段、提示词和任务载荷等信息。",
    ),
    (
        "5",
        "asset_images",
        "图像资产表",
        "用于保存任务输出的图像名称、文件路径、所属模型和生成状态等资产信息。",
    ),
    (
        "6",
        "asset_image_prompts",
        "图像提示词参数表",
        "用于记录单张图像对应的提示词内容、采样步数、随机种子和引导系数等生成参数。",
    ),
    (
        "7",
        "asset_image_scores",
        "图像评分结果表",
        "用于保存图像在视觉保真度、文本一致性、物理合理性和构图美观度等维度上的评分结果。",
    ),
    (
        "8",
        "asset_image_score_explanations",
        "图像评分解释表",
        "用于保存图像评分解释结果、检查后图像路径以及结构化解释内容。",
    ),
    (
        "9",
        "audit_task_events",
        "任务审计事件表",
        "用于记录任务执行过程中的事件类型、消息内容和扩展载荷，实现任务审计追踪。",
    ),
]


def set_run_font(run, east_asia_font: str, latin_font: str, size_pt: float, *, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = latin_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)


def set_paragraph_style(paragraph, *, align: int, size_pt: float, bold: bool = False, first_line_indent_pt: float = 0) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", size_pt, bold=bold)


def set_cell_borders(cell, *, top: str = "nil", bottom: str = "nil", left: str = "nil", right: str = "nil") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, value in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), value)
        if value != "nil":
            element.set(qn("w:sz"), "8")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")


def fill_cell(cell, text: str, *, align: int, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.add_run(text)
    set_paragraph_style(paragraph, align=align, size_pt=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title = doc.add_paragraph()
    title.add_run("Electric AI Platform 数据库简表")
    set_paragraph_style(title, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=18, bold=True)

    doc.add_paragraph()

    caption = doc.add_paragraph()
    caption.add_run("表3.1 数据库主要数据表设计")
    set_paragraph_style(caption, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10.5)

    table = doc.add_table(rows=1 + len(ROWS), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(1.5), Cm(4.0), Cm(3.0), Cm(7.0)]
    headers = ["序号", "表名", "中文名称", "简要说明"]

    for column, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[column]
        cell.width = width
        fill_cell(cell, header, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        set_cell_borders(cell, top="single", bottom="single")

    for row_index, values in enumerate(ROWS, start=1):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.85)
        for column, (text, width) in enumerate(zip(values, widths)):
            cell = row.cells[column]
            cell.width = width
            align = WD_ALIGN_PARAGRAPH.LEFT if column == 3 else WD_ALIGN_PARAGRAPH.CENTER
            fill_cell(cell, text, align=align, bold=False)
            is_last_row = row_index == len(ROWS)
            set_cell_borders(cell, bottom="single" if is_last_row else "nil")

    return doc


def main() -> int:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
